import docx

from app.parsing.base import ParsedDocument


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    document = docx.Document(file_bytes)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    return ParsedDocument(
        source_filename=filename,
        raw_text="\n".join(paragraphs),
        pages=None,
        extraction_method="native",
    )
